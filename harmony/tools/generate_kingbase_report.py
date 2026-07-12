# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SmartSite_Kingbase_Database_Report.docx"
ASSET_DIR = ROOT / "docs" / "kingbase_report_assets"
REQUIREMENT_IMAGE = ASSET_DIR / "kingbase_requirement.png"
API_START_IMAGE = ASSET_DIR / "kingbase_api_start.png"


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("智慧工地系统国产数据库应用说明  |  ")
    set_font(run, 9, False, "6B7280")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    field.append(r)
    footer._p.append(field)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def write_cell(cell, text, bold=False, color="1F2937"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_font(run, 10, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], "F2F4F7")
        write_cell(table.rows[0].cells[i], header, True, "0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            write_cell(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph("")
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("• " + text)
    set_font(run, 11)


def add_number(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{num}. {text}")
    set_font(run, 11)


def add_image(doc, image_path, caption, width=6.2):
    if not image_path.exists():
        return False
    doc.add_picture(str(image_path), width=Inches(width))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.runs[0], 9, False, "6B7280")
    return True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("智慧工地系统国产数据库（金仓数据库）应用说明")
    set_font(run, 22, True, "0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于 KingbaseES 的材料监测历史数据存储、接口服务与鸿蒙端调用总结")
    set_font(run, 11, False, "4B5563")


def build_report():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    add_table(
        doc,
        ["项目项", "说明"],
        [
            ["项目名称", "智构云工地管理系统（智慧工地项目）"],
            ["国产数据库", "金仓数据库 KingbaseES"],
            ["数据库用途", "保存材料区监测历史数据、区域配置、告警规则、告警事件和数字孪生资产信息"],
            ["后端接口", "D:\\SmartSiteServices\\kingbase-material-api"],
            ["鸿蒙端配置", "entry/src/main/ets/common/MaterialDatabaseConfig.ets"],
            ["报告定位", "说明本项目如何使用国产数据库完成材料监测相关业务功能"],
        ],
        [1.5, 4.9],
    )

    doc.add_heading("1. 任务要求与项目对应关系", level=1)
    doc.add_paragraph(
        "课程任务要求基于国产数据库（金仓数据库）开发项目，在金仓数据库上建立 3 个以上数据表，并在 Web 端或移动终端至少完成一个与项目相关的业务功能。智慧工地项目将金仓数据库用于材料监测模块，围绕水泥区和木材区的传感器数据进行历史保存、查询展示和趋势分析。"
    )
    add_image(doc, REQUIREMENT_IMAGE, "图 1  国产数据库课程任务要求", 5.7)
    add_table(
        doc,
        ["要求项", "项目实现情况"],
        [
            ["使用国产数据库", "使用金仓数据库 KingbaseES 保存智慧工地材料区数据。"],
            ["建立 3 个以上数据表", "已设计 users、material_zones、material_sensor_records、material_alert_rules、material_alert_events、digital_twin_assets 共 6 张表。"],
            ["完成业务功能", "鸿蒙端材料区可将华为云 IoTDA 数据写入金仓，并按水泥区、木材区查询历史记录绘制折线趋势图。"],
            ["移动终端使用", "HarmonyOS 应用通过 HTTP 调用金仓材料 API，移动端不直接保存数据库账号密码。"],
        ],
        [1.8, 4.6],
    )

    doc.add_heading("2. 金仓数据库在系统中的定位", level=1)
    doc.add_paragraph(
        "智慧工地系统中存在多类数据来源：华为云 IoTDA 负责实时设备影子和命令下发，Supabase 云数据库负责打卡任务和跨设备签到记录，金仓数据库负责材料监测历史数据沉淀。这样的分工可以让每类数据进入最合适的存储位置。"
    )
    add_table(
        doc,
        ["数据模块", "使用的数据平台", "原因"],
        [
            ["实时设备状态", "华为云 IoTDA", "设备端上报温湿度、灯、风扇等实时状态，适合通过设备影子读取。"],
            ["打卡任务与记录", "Supabase 云数据库", "需要多台手机和管理员端共享任务与记录，适合云端数据库。"],
            ["材料历史数据", "金仓数据库 KingbaseES", "满足国产数据库课程要求，并保存温湿度等历史记录用于折线趋势和告警分析。"],
            ["数字孪生配置", "金仓数据库 KingbaseES", "保存材料区、传感器、风扇等孪生资产元数据。"],
        ],
        [1.6, 2.0, 2.8],
    )

    doc.add_heading("3. 数据流与系统架构", level=1)
    doc.add_paragraph(
        "金仓数据库没有直接暴露给鸿蒙手机端。项目采用“鸿蒙端 + Node.js 后端接口 + 金仓数据库”的结构，避免在移动端保存数据库密码，也避免手机端直接安装数据库驱动。"
    )
    add_table(
        doc,
        ["步骤", "数据流说明"],
        [
            ["1", "板端设备将温度、湿度、灯、风扇等数据上报到华为云 IoTDA。"],
            ["2", "HarmonyOS 应用定时读取华为云 IoTDA 设备影子，拿到水泥区和木材区最新数据。"],
            ["3", "鸿蒙端通过 POST /api/material-records 将材料区数据发送到 Kingbase Material API。"],
            ["4", "Node.js API 使用 pg 驱动连接 KingbaseES，将记录写入 material_sensor_records。"],
            ["5", "后端根据 material_alert_rules 判断是否生成 material_alert_events 告警事件。"],
            ["6", "用户点击材料区历史趋势时，鸿蒙端通过 GET /api/material-records 查询金仓数据并绘制折线图。"],
        ],
        [0.7, 5.7],
    )

    doc.add_heading("4. 数据库表设计", level=1)
    doc.add_paragraph(
        "本项目在金仓数据库中建立了 6 张核心业务表，超过课程要求的 3 张表。表结构覆盖用户、材料区配置、传感器历史记录、告警规则、告警事件和数字孪生资产。"
    )
    add_table(
        doc,
        ["表名", "主要作用", "关键字段"],
        [
            ["users", "保存材料管理员、材料区负责人等本地管理用户。", "username、real_name、role、phone、remark"],
            ["material_zones", "保存水泥区、木材区等材料区配置和存储阈值。", "zone_code、zone_name、manager_name、temp_min、temp_max、humi_min、humi_max"],
            ["material_sensor_records", "保存华为云 IoTDA 同步过来的材料区传感器历史数据。", "zone_code、device_id、temp、humi、dist、lumi、lamp_status、created_at"],
            ["material_alert_rules", "保存温度、湿度等指标的告警规则。", "zone_code、metric_code、min_value、max_value、severity、enabled"],
            ["material_alert_events", "保存实际触发的超温、超湿等告警事件。", "record_id、zone_code、metric_code、actual_value、severity、status、message"],
            ["digital_twin_assets", "保存数字孪生场景中的区域、风扇、传感器等对象信息。", "asset_code、asset_name、asset_type、zone_code、model_type、status"],
        ],
        [1.6, 2.3, 2.5],
    )

    doc.add_heading("5. 后端接口服务实现", level=1)
    doc.add_paragraph(
        "金仓材料接口位于 D:\\SmartSiteServices\\kingbase-material-api。该服务使用 Node.js 编写，通过 pg 依赖连接 KingbaseES。项目采用 8088 端口向鸿蒙端提供 HTTP 接口。"
    )
    add_image(doc, API_START_IMAGE, "图 2  Kingbase Material API 启动成功截图", 6.2)
    add_table(
        doc,
        ["接口", "方法", "作用"],
        [
            ["/api/health", "GET", "检测接口服务和金仓数据库连接状态。"],
            ["/api/material-records", "POST", "新增材料区传感器历史记录，并根据告警规则生成告警事件。"],
            ["/api/material-records?zoneCode=...&limit=...", "GET", "按材料区查询最近若干条历史记录，用于趋势图。"],
            ["/api/material-zones", "GET", "查询水泥区、木材区等材料区配置与负责人信息。"],
            ["/api/material-alerts?zoneCode=...&limit=...", "GET", "查询材料区告警事件。"],
            ["/api/twin-assets", "GET", "查询数字孪生资产配置。"],
        ],
        [2.7, 0.8, 2.9],
    )

    doc.add_heading("6. 数据库初始化与服务启动过程", level=1)
    doc.add_paragraph(
        "金仓数据库服务需要先安装并创建数据库，然后执行 schema.sql 建表。后端服务通过 .env 文件读取数据库地址、端口、数据库名、用户名和密码。"
    )
    for idx, item in enumerate(
        [
            "在金仓数据库中创建 smart_site 数据库。",
            "进入 D:\\SmartSiteServices\\kingbase-material-api 目录。",
            "根据 .env.example 创建 .env，填写 KINGBASE_HOST、KINGBASE_PORT、KINGBASE_DATABASE、KINGBASE_USER、KINGBASE_PASSWORD 等连接信息。",
            "执行 npm install 安装 Node.js 依赖。",
            "执行 npm run init-db 初始化表结构和基础数据。",
            "执行 cmd /c npm start 启动服务，看到 Kingbase material API listening on http://0.0.0.0:8088 表示服务启动。",
            "在鸿蒙端 MaterialDatabaseConfig.ets 中配置 apiBaseUrl，例如 http://电脑IP:8088。",
        ],
        1,
    ):
        add_number(doc, idx, item)

    doc.add_heading("7. 鸿蒙端如何使用金仓数据库", level=1)
    doc.add_paragraph(
        "鸿蒙端不直接连接金仓数据库，而是通过 Kingbase Material API 访问。项目中的配置文件为 entry/src/main/ets/common/MaterialDatabaseConfig.ets，当前配置包含 apiBaseUrl、zoneCode、zoneName 和 requestTimeoutMs。"
    )
    add_table(
        doc,
        ["鸿蒙端逻辑", "实现说明"],
        [
            ["材料数据写入", "refreshCloudData 从华为云 IoTDA 获取设备影子后，调用 saveMaterialSensorRecord 将水泥区和木材区数据写入金仓接口。"],
            ["历史数据查询", "用户点击水泥区或木材区卡片后，openMaterialHistory 设置 zoneCode，再调用 loadMaterialHistory。"],
            ["趋势图绘制", "loadMaterialHistoryAsync 请求 /api/material-records，拿到 records 后反转为时间顺序并调用 drawMaterialTrendChart。"],
            ["异常提示", "如果接口未启动或数据库不可用，页面显示“请启动金仓材料 API 并创建表”等提示。"],
            ["数字孪生数据", "材料孪生页面可读取 material-zones、material-alerts、twin-assets 和最近一条 material-records 数据。"],
        ],
        [1.6, 4.8],
    )

    doc.add_heading("8. 与材料监测业务的结合", level=1)
    doc.add_paragraph(
        "金仓数据库在本项目中不是孤立的数据库练习，而是直接服务材料监测业务。管理员在材料区看到实时温湿度后，可以点击水泥区或木材区进入历史趋势页面，查看一段时间内温度和湿度变化。"
    )
    add_table(
        doc,
        ["业务场景", "金仓数据库提供的能力"],
        [
            ["水泥区历史趋势", "按 cement_zone 查询历史记录，显示温度、湿度曲线。"],
            ["木材区历史趋势", "按 timber_zone 查询历史记录，比较木材存储环境变化。"],
            ["材料告警", "根据 material_alert_rules 判断超温、超湿，并写入 material_alert_events。"],
            ["负责人信息", "material_zones 关联 users，保存每个材料区负责人和联系电话。"],
            ["数字孪生联动", "digital_twin_assets 保存虚拟区域、传感器、风扇等对象，用于 3D 场景数据绑定。"],
        ],
        [1.7, 4.7],
    )

    doc.add_heading("9. 增删改查完成情况说明", level=1)
    doc.add_paragraph(
        "从当前项目代码看，金仓材料模块已经完成与项目业务直接相关的“新增”和“查询”：新增传感器历史记录，查询历史记录、区域配置、告警事件和数字孪生资产。课程要求中的“增删改查”如果按完整 CRUD 严格验收，后续可以在 material_zones、material_alert_rules 或 material_alert_events 上补充修改和删除接口。"
    )
    add_table(
        doc,
        ["操作类型", "当前项目体现", "说明"],
        [
            ["新增 Create", "POST /api/material-records", "鸿蒙端把华为云 IoTDA 数据写入金仓数据库。"],
            ["查询 Read", "GET /api/material-records、/api/material-zones、/api/material-alerts、/api/twin-assets", "鸿蒙端读取历史趋势、区域配置、告警和孪生数据。"],
            ["修改 Update", "表结构支持规则、区域、告警状态维护", "当前移动端主要聚焦监测展示，完整修改接口可作为后续扩展。"],
            ["删除 Delete", "可针对历史记录或规则维护扩展删除接口", "当前项目未把删除作为材料监测主流程，避免误删演示数据。"],
        ],
        [1.2, 2.6, 2.6],
    )

    doc.add_heading("10. 技术价值总结", level=1)
    for item in [
        "满足国产数据库课程要求：项目使用 KingbaseES，并建立 6 张业务表，超过 3 张表要求。",
        "与智慧工地业务强相关：金仓数据库保存的是材料监测历史、告警和数字孪生资产，而不是脱离项目的示例表。",
        "移动端可实际调用：HarmonyOS 应用通过 HTTP 接口读取金仓数据，并在材料区历史趋势页面展示。",
        "数据来源真实：历史数据来自华为云 IoTDA 设备影子同步，不在报告和界面中伪造业务数据。",
        "架构更安全：手机端不直接保存数据库密码，数据库访问集中在 Node.js 后端接口中。",
        "便于后续扩展：表结构已经包含告警规则、告警事件、区域负责人和孪生资产，后续可继续扩展完整 CRUD 管理页面。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("11. 总结", level=1)
    doc.add_paragraph(
        "智慧工地项目将国产数据库 KingbaseES 应用于材料监测历史数据管理。系统通过 Node.js 后端接口连接金仓数据库，鸿蒙端通过 HTTP 调用接口完成材料记录写入、历史记录查询、趋势图展示、告警读取和数字孪生数据读取。"
    )
    doc.add_paragraph(
        "从整体实现看，金仓数据库承担了项目中“历史沉淀”和“分析支撑”的角色，使实时 IoTDA 数据能够被保存、回看和用于趋势分析。该模块既满足国产数据库建表与业务功能要求，也与智慧工地的材料区温湿度监测场景保持一致，具有明确的项目应用价值。"
    )

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_font(run)

    doc.core_properties.title = "智慧工地系统国产数据库（金仓数据库）应用说明"
    doc.core_properties.subject = "KingbaseES 智慧工地材料监测数据库应用总结"
    doc.core_properties.author = "Codex"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
