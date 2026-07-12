# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SmartSite_LLM_Tech_Summary_Report.docx"
ASSET_DIR = ROOT / "docs" / "report_assets"
ASSISTANT_UI_IMAGE = ASSET_DIR / "harmony_ai_assistant.png"
WORKFLOW_IMAGE = ASSET_DIR / "dify_workflow.png"
OLLAMA_PROVIDER_IMAGE = ASSET_DIR / "ollama_provider.png"


def set_east_asia_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("智慧工地系统大模型技术应用总结报告  |  ")
    set_east_asia_font(run, 9, False, "6B7280")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    field.append(r)
    footer._p.append(field)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def write_cell(cell, text, bold=False, color="1F2937"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_east_asia_font(run, 10, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], "F2F4F7")
        write_cell(table.rows[0].cells[i], header, True, "0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            write_cell(cells[i], str(value))
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph("")
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("• " + text)
    set_east_asia_font(run, 11)


def add_number(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{num}. {text}")
    set_east_asia_font(run, 11)


def add_image_with_caption(doc, image_path, caption, width=6.3):
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_east_asia_font(cap.runs[0], 9, False, "6B7280")
        return True
    return False


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("智慧工地系统大模型技术应用总结报告")
    set_east_asia_font(run, 22, True, "0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("面向 HarmonyOS 智慧工地项目的 AI 助手、数据检索与业务编排说明")
    set_east_asia_font(run, 11, False, "4B5563")


def build_report():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    add_table(
        doc,
        ["项目项", "说明"],
        [
            ["项目名称", "智构云工地管理系统（智慧工地项目）"],
            ["应用端", "HarmonyOS 鸿蒙端，包含工人端与管理员端"],
            ["大模型入口", "应用内“AI / 智慧工地助手”悬浮入口"],
            ["报告定位", "总结系统中已经接入并使用的大模型相关技术，不作为优化建议书"],
        ],
        [1.6, 4.8],
    )

    doc.add_heading("1. 报告目的", level=1)
    doc.add_paragraph(
        "本报告用于说明“智慧工地”系统在实际开发中如何接入和使用大模型技术。报告重点描述已经在系统中形成的功能、数据链路和业务作用，便于答辩、项目验收和后续维护时说明大模型在本项目中的位置。"
    )
    doc.add_paragraph(
        "系统中的大模型能力不是孤立的聊天窗口，而是与打卡系统、材料监测、华为云 IoTDA、金仓数据库、Supabase 云数据库、THINGJS 数字孪生和项目知识库共同组成的智能交互层。用户可以通过自然语言查询项目功能、打卡状态、材料区数据、监控接入方式和运维问题。"
    )

    doc.add_heading("2. 项目已实现功能概述", level=1)
    doc.add_paragraph(
        "智慧工地系统当前已经形成了面向工人端和管理员端的应用结构，并接入了多类真实数据源。项目的核心功能如下："
    )
    for item in [
        "工人端与管理员端分离：工人端聚焦签到、签退和秒级工时显示，管理员端聚焦看板、材料监测、打卡管理和设备控制。",
        "打卡系统：管理员可以管理打卡任务，工人端可以进行签到和签退，打卡记录继续使用 Supabase 云数据库保存，便于不同设备共享记录。",
        "材料监测：水泥区与木材区分别接入华为云 IoTDA 设备影子，读取温度、湿度、灯、风扇等状态，并在管理端材料区展示。",
        "命令下发：材料区风扇旁提供 ON/OFF 控制按钮，水泥区通过 SetFan/FanSt 下发，木材区通过 Fan/fan 下发，用于控制板端设备。",
        "金仓数据库：本地金仓数据库用于保存材料监测历史数据，材料区点击后可查看对应区域的折线趋势图。",
        "视频监控：系统保留 RTSP 视频流接入入口，用于展示施工现场摄像头画面。",
        "数字孪生：系统通过 WebView 接入 THINGJS 场景，用于展示材料区和工地空间的 3D 可视化。",
        "智慧工地助手：工人端和管理员端均提供 AI 助手入口，用户可以用自然语言查询打卡、材料、视频、孪生和项目运维问题。",
    ]:
        add_bullet(doc, item)

    add_table(
        doc,
        ["功能模块", "主要数据来源", "在系统中的作用"],
        [
            ["打卡签到", "Supabase 云数据库", "保存任务、记录工人签到签退和工时数据"],
            ["材料监测", "华为云 IoTDA 设备影子", "读取水泥区、木材区实时传感器数据"],
            ["历史趋势", "金仓数据库", "保存材料区历史记录并绘制折线图"],
            ["设备控制", "华为云 IoTDA 命令下发", "向水泥区、木材区板端下发风扇控制命令"],
            ["数字孪生", "THINGJS 场景链接", "展示工地和材料区 3D 场景"],
            ["AI 助手", "Dify + Ollama/DeepSeek + 知识库 + HTTP 节点", "理解用户问题，调用数据并生成自然语言回答"],
        ],
        [1.4, 2.0, 3.0],
    )

    doc.add_heading("3. 大模型技术总体架构", level=1)
    doc.add_paragraph(
        "本项目的大模型技术采用“鸿蒙端入口 + Dify 编排 + 本地模型推理 + 业务数据查询”的结构。鸿蒙应用负责显示聊天界面和发送用户问题；Dify 负责问题分类、数据调用、知识检索和节点编排；Ollama 提供本地大模型运行能力；Supabase、金仓数据库和华为云 IoTDA 提供真实业务数据。"
    )
    add_table(
        doc,
        ["层级", "组件", "职责"],
        [
            ["交互层", "HarmonyOS 智慧工地助手页面", "接收用户问题，展示大模型回答"],
            ["编排层", "Dify 工作流", "进行问题分类、HTTP 请求、知识检索、代码整理和回复编排"],
            ["模型层", "Ollama / DeepSeek 模型", "完成自然语言理解、归纳和回答生成"],
            ["业务数据层", "Supabase、金仓数据库、IoTDA 接口", "提供打卡任务、材料历史数据和设备状态"],
            ["展示层", "看板、材料区、折线图、THINGJS 场景", "将业务数据和 AI 说明呈现给用户"],
        ],
        [1.2, 2.1, 3.1],
    )

    doc.add_heading("4. 大模型部署与接入过程", level=1)
    doc.add_paragraph(
        "本系统的大模型能力采用本地模型服务与 Dify 工作流编排相结合的方式部署。部署过程重点解决三个问题：模型在哪里运行、业务数据如何进入模型、模型回答如何回到鸿蒙端。"
    )

    doc.add_heading("4.1 本地模型服务部署", level=2)
    doc.add_paragraph(
        "本项目使用 Ollama 作为本地大模型运行服务，模型选择 DeepSeek 系列。Ollama 负责在本机启动模型推理接口，Dify 作为上层应用平台调用该模型。"
    )
    add_table(
        doc,
        ["步骤", "操作内容", "作用"],
        [
            ["安装 Ollama", "在电脑端安装 Ollama，并将模型目录按实际磁盘空间配置到 D 盘。", "提供本地大模型运行环境"],
            ["拉取模型", "执行 ollama pull deepseek-r1:7b。", "下载项目使用的 DeepSeek 模型"],
            ["启动服务", "执行 ollama serve，保持本地模型服务运行。", "向 Dify 提供模型推理接口"],
            ["验证模型", "通过 ollama list 或在 Dify 模型供应商页面查看模型。", "确认模型已经可被调用"],
        ],
        [1.25, 2.75, 2.4],
    )
    add_image_with_caption(
        doc,
        OLLAMA_PROVIDER_IMAGE,
        "图 1  Dify 中配置 Ollama 模型供应商与 deepseek-r1:7b 模型",
        6.3,
    )

    doc.add_heading("4.2 Dify 应用与工作流配置", level=2)
    doc.add_paragraph(
        "Dify 用于把本地大模型包装成“智慧工地助手”。在 Dify 中创建应用后，需要配置系统提示词、问题分类器、HTTP 请求节点、知识检索节点、代码执行节点和 LLM 节点。这样用户问题进入系统后，不会直接交给模型泛泛回答，而是先按业务类型调用真实数据。"
    )
    add_table(
        doc,
        ["配置项", "本项目中的配置方式"],
        [
            ["模型供应商", "选择 Ollama，并绑定 deepseek-r1:7b 作为主要对话模型。"],
            ["系统提示词", "设定助手身份为“智慧工地助手”，要求围绕打卡、材料监测、视频监控、数字孪生和项目运维回答。"],
            ["知识库", "导入项目说明文档和使用说明，作为回答项目功能、操作流程和排查步骤的依据。"],
            ["HTTP 节点", "分别访问材料监测接口、木材区接口、水泥区接口和 Supabase 打卡任务接口。"],
            ["代码执行节点", "整理接口返回的 JSON 数据，提取温度、湿度、任务状态等关键字段。"],
            ["直接回复节点", "把最终 answer 返回给鸿蒙端，由应用页面显示给用户。"],
        ],
        [1.6, 4.8],
    )

    doc.add_heading("4.3 鸿蒙端调用 Dify 接口", level=2)
    doc.add_paragraph(
        "鸿蒙端在工人端和管理员端都提供 AI 悬浮入口。用户点击后进入“智慧工地助手”页面，页面通过 HTTP 请求访问 Dify 的 chat-messages 接口，并把返回结果显示在聊天区域。"
    )
    add_table(
        doc,
        ["调用项", "说明"],
        [
            ["请求地址", "本机调试可使用 http://localhost/v1/chat-messages；真机访问时需要换成电脑局域网 IP 或内网穿透地址。"],
            ["请求方法", "POST"],
            ["鉴权方式", "Authorization: Bearer app-应用密钥。报告中不写入真实密钥，避免泄露。"],
            ["请求体", "包含 inputs、query、response_mode、conversation_id 和 user 等字段。"],
            ["返回字段", "鸿蒙端主要读取 answer 字段并显示给用户。"],
        ],
        [1.4, 5.0],
    )
    add_image_with_caption(
        doc,
        ASSISTANT_UI_IMAGE,
        "图 2  鸿蒙端智慧工地助手页面",
        3.0,
    )

    doc.add_heading("4.4 运行顺序", level=2)
    doc.add_paragraph(
        "为了保证助手能够查到真实数据，演示时需要按依赖顺序启动相关服务。基础数据服务先启动，大模型和 Dify 后启动，最后运行鸿蒙应用。"
    )
    for idx, item in enumerate(
        [
            "启动金仓数据库和材料数据 API，保证材料历史数据接口可访问。",
            "确认 Supabase 云数据库可访问，打卡任务和签到记录可以被 HTTP 节点查询。",
            "确认华为云 IoTDA 设备影子和命令下发接口可用，材料区实时数据能够同步。",
            "启动 Ollama 服务，并确认 deepseek-r1:7b 模型存在。",
            "启动 Dify，本地模型供应商、知识库和工作流节点保持可用。",
            "运行 HarmonyOS 智慧工地应用，从 AI 入口进入助手页面进行问答。",
        ],
        1,
    ):
        add_number(doc, idx, item)

    doc.add_heading("5. Dify 工作流实现说明", level=1)
    doc.add_paragraph(
        "根据当前 Dify 编排界面，智慧工地助手已经配置为多节点工作流。工作流不是单纯把用户输入直接交给模型，而是先判断问题类型，再按类型调用相应数据源或知识库，最后由大模型生成更符合业务语境的回答。"
    )
    add_image_with_caption(
        doc,
        WORKFLOW_IMAGE,
        "图 3  智慧工地助手 Dify 工作流编排界面",
        6.3,
    )

    add_table(
        doc,
        ["节点类型", "当前用途"],
        [
            ["用户输入", "接收工人端或管理员端发出的自然语言问题。"],
            ["问题分类器", "将问题分为材料监测、打卡签到、监控视频、数字孪生、项目介绍和其他问题。"],
            ["HTTP 请求", "访问材料区接口，查询水泥区和木材区最近一条监测数据。"],
            ["HTTP 请求 3", "访问 Supabase 接口，查询当前有效打卡任务。"],
            ["知识检索", "检索项目说明文档、使用说明和业务规则。"],
            ["代码执行", "对 HTTP 返回结果进行整理，提取温度、湿度、任务状态等关键字段。"],
            ["LLM 节点", "根据用户问题、分类结果、实时数据和知识库内容生成回答。"],
            ["直接回复", "将最终回答返回给鸿蒙端聊天界面。"],
        ],
        [1.6, 4.8],
    )

    for idx, item in enumerate(
        [
            "用户在鸿蒙端点击 AI 悬浮按钮，进入智慧工地助手页面。",
            "用户输入问题，例如“水泥区温度”“今天有没有打卡任务”“RTSP 视频为什么加载失败”。",
            "Dify 问题分类器判断问题所属业务模块。",
            "若问题需要实时数据，Dify 通过 HTTP 节点访问对应接口，例如材料监测接口或 Supabase 打卡接口。",
            "代码执行节点整理接口返回值，避免把原始 JSON 直接丢给用户。",
            "LLM 节点结合用户问题、业务数据和知识库内容生成自然语言说明。",
            "鸿蒙端展示 Dify 返回的 answer 字段，形成完整问答体验。",
        ],
        1,
    ):
        add_number(doc, idx, item)

    doc.add_heading("6. 大模型在业务中的具体作用", level=1)
    sections = [
        (
            "6.1 打卡签到问答",
            "打卡系统已经接入 Supabase 云数据库。大模型助手可以围绕当前打卡任务、签到签退逻辑、工时统计规则、管理员查看记录等问题进行解释。对于工人端用户，它能说明当前是否需要签到、为什么需要签退以及工时如何计算；对于管理员端用户，它能说明如何发布任务、查看记录和理解打卡数据。",
        ),
        (
            "6.2 材料监测问答",
            "材料区数据来自华为云 IoTDA 和金仓数据库。大模型助手可以结合水泥区、木材区的温度、湿度、风扇状态、灯状态和历史趋势，回答“当前温湿度是否正常”“为什么风扇没有打开”“历史趋势图代表什么”等问题。对于接口返回为空、数据未同步或设备未上报的情况，助手应明确说明当前没有取到数据，而不是编造数值。",
        ),
        (
            "6.3 设备命令下发说明",
            "系统已经在材料监测区加入风扇控制按钮。水泥区和木材区使用不同的华为云产品模型命令：水泥区使用 service_id 为 123 的 SetFan 命令，参数为 FanSt；木材区使用 Fan 命令，参数为 fan。大模型助手可以解释命令下发成功、设备执行、设备重新上报影子三者之间的关系，帮助用户理解“云端显示已下发但界面状态未变化”的原因。",
        ),
        (
            "6.4 视频监控与 RTSP 接入说明",
            "系统保留视频监控入口，用于接入 RTSP 视频流。大模型助手可以根据知识库说明，解释本机地址、局域网地址、公网地址之间的区别，也可以说明视频一直加载可能与推流服务、播放地址、网络连通性或端口开放有关。",
        ),
        (
            "6.5 数字孪生说明",
            "系统通过 WebView 接入 THINGJS 分享场景，用于展示工地和材料区的三维空间。大模型助手可以向用户解释数字孪生页面的用途、材料区设备命名规则、温湿度数据与 3D 场景之间的关系，以及为什么需要通过接口把真实数据绑定到场景表现上。",
        ),
    ]
    for heading, text in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(text)

    doc.add_heading("7. 数据真实性与回答边界", level=1)
    doc.add_paragraph(
        "本项目明确要求数据不能作假。因此，大模型在系统中的定位是“解释和组织真实数据”，不是生成虚构数据。工作流中的 HTTP 节点、知识检索节点和代码执行节点承担了数据获取与整理职责，大模型只在已有数据和文档的基础上生成表达更自然的回答。"
    )
    add_table(
        doc,
        ["数据类型", "真实来源", "大模型处理方式"],
        [
            ["打卡任务与记录", "Supabase 云数据库", "读取任务状态和记录后，用自然语言解释给用户"],
            ["实时温湿度", "华为云 IoTDA 设备影子", "根据接口返回值说明当前材料区状态"],
            ["历史趋势", "金仓数据库", "根据历史记录解释曲线变化和区间状态"],
            ["设备控制结果", "华为云 IoTDA 命令下发与设备影子", "说明命令已下发、设备是否上报、状态为何未同步"],
            ["项目规则说明", "知识库文档", "按项目文档回答操作流程和功能说明"],
        ],
        [1.6, 2.0, 2.8],
    )
    doc.add_paragraph(
        "如果接口访问失败、数据库没有记录或设备没有上报，助手应返回“当前未获取到数据”或“连接失败，请检查接口、网络或 Token”，这也是系统保持数据真实性的一部分。"
    )

    doc.add_heading("8. 系统价值总结", level=1)
    for item in [
        "降低使用门槛：工人和管理员不需要理解全部后台结构，也可以通过自然语言询问系统状态。",
        "提高演示完整度：大模型助手把打卡、材料监测、设备控制、数字孪生和视频监控串成一个可说明的整体。",
        "增强运维解释能力：当 Token 失效、接口 403、命令已下发但影子未变化等问题出现时，助手可以解释排查方向。",
        "体现国产数据库与云端数据结合：金仓数据库负责材料历史数据，Supabase 保留跨设备打卡记录，华为云 IoTDA 负责设备接入。",
        "为后续智能预警打基础：现有数据链路已经具备实时读取、历史记录和自然语言解释能力。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("9. 未来展望", level=1)
    doc.add_paragraph(
        "以下内容属于后续可继续拓展的方向，不影响当前系统已经完成的大模型助手、数据查询和业务说明能力。"
    )
    for item in [
        "更丰富的数字孪生联动：在 THINGJS 场景中进一步绑定风扇旋转、灯光亮灭、湿度告警颜色和设备点击标签。",
        "更细的异常分析：基于金仓数据库中的历史数据，加入温湿度异常趋势判断和更清晰的报警解释。",
        "更安全的密钥管理：演示阶段使用本地配置较多，后续正式部署可将 Token、API Key 等敏感信息迁移到后端服务管理。",
        "多模态监控分析：在 RTSP 视频接入稳定后，可进一步探索施工现场画面识别、人员安全帽识别等能力。",
        "语音交互：在移动端加入语音输入和语音播报，让工地现场查询更加方便。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("10. 总结", level=1)
    doc.add_paragraph(
        "智慧工地系统已经把大模型技术落到具体业务场景中：用户从鸿蒙端发起问题，Dify 工作流根据问题类型调用材料监测、打卡任务、知识库等数据源，Ollama/DeepSeek 模型负责理解和组织回答，最终在应用内形成智慧工地助手体验。"
    )
    doc.add_paragraph(
        "从项目整体看，大模型承担的是“智能交互入口”和“业务数据解释器”的角色。它连接了工人端、管理员端、华为云 IoTDA、Supabase、金仓数据库、RTSP 视频和 THINGJS 数字孪生，使系统不仅能展示数据，也能用自然语言解释数据、说明操作和辅助排查问题。这体现了大模型技术在智慧工地管理系统中的实际应用价值。"
    )

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_east_asia_font(run)

    doc.core_properties.title = "智慧工地系统大模型技术应用总结报告"
    doc.core_properties.subject = "HarmonyOS 智慧工地项目大模型技术应用总结"
    doc.core_properties.author = "Codex"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
