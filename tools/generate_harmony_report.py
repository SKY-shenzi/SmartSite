# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SmartSite_HarmonyOS_Tech_Report.docx"
ASSET_DIR = ROOT / "docs" / "harmony_report_assets"
APP_HOME_IMAGE = ASSET_DIR / "app_runtime_home.png"
PROJECT_STRUCTURE_IMAGE = ASSET_DIR / "deveco_project_structure.png"
BUILD_SUCCESS_IMAGE = ASSET_DIR / "deveco_build_success.png"


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("智慧工地系统鸿蒙开发技术应用说明  |  ")
    set_font(run, 9, False, "6B7280")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    field.append(r)
    footer._p.append(field)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def write_cell(cell, text, bold=False, color="1F2937"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_font(run, 10, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], "F2F4F7")
        write_cell(table.rows[0].cells[i], header, True, "0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            write_cell(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph("")
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("• " + text)
    set_font(run, 11)


def add_number(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{num}. {text}")
    set_font(run, 11)


def add_image(doc, image_path, caption, width=6.2):
    if not image_path.exists():
        return False
    doc.add_picture(str(image_path), width=Inches(width))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.runs[0], 9, False, "6B7280")
    return True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("智慧工地系统鸿蒙开发技术应用说明")
    set_font(run, 22, True, "0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于 HarmonyOS、ArkTS、ArkUI 与 DevEco Studio 的项目实现总结")
    set_font(run, 11, False, "4B5563")


def build_report():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    add_table(
        doc,
        ["项目项", "说明"],
        [
            ["项目名称", "智构云工地管理系统（智慧工地项目）"],
            ["开发平台", "DevEco Studio"],
            ["运行平台", "HarmonyOS 手机端"],
            ["开发语言与 UI", "ArkTS + ArkUI 声明式开发"],
            ["工程类型", "entry 模块应用工程，包含 UIAbility 与 FormExtensionAbility"],
            ["报告定位", "说明本项目如何实际使用鸿蒙开发技术完成应用功能"],
        ],
        [1.5, 4.9],
    )

    doc.add_heading("1. 报告目的", level=1)
    doc.add_paragraph(
        "本报告用于说明智慧工地系统在开发过程中如何使用鸿蒙开发技术。报告不是单纯介绍 HarmonyOS 概念，而是结合当前项目的工程结构、页面实现、网络访问、桌面卡片、构建运行和业务功能，说明鸿蒙技术在项目中的实际落地方式。"
    )
    doc.add_paragraph(
        "项目面向施工现场管理场景，包含管理员端、工人端、材料监测、打卡、视频监控、数字孪生、AI 助手和桌面卡片等功能。鸿蒙开发技术承担了移动端界面展示、设备能力调用、网络接口访问、状态管理和应用打包部署等核心工作。"
    )

    doc.add_heading("2. 鸿蒙应用形态与本项目定位", level=1)
    doc.add_paragraph(
        "HarmonyOS 应用通常包括传统安装式应用、元服务和万能卡片等形态。本项目当前以 HarmonyOS 应用工程为主体，面向手机端运行，同时配置了桌面卡片能力，用于展示智慧工地关键数据。"
    )
    add_table(
        doc,
        ["鸿蒙形态", "含义", "本项目使用情况"],
        [
            ["HarmonyOS 应用", "安装后在手机、平板等设备上运行的完整应用。", "项目主体采用该形态，提供登录、看板、材料、监控、打卡、工时等页面。"],
            ["元服务", "轻量化、服务直达、免安装体验的应用形态。", "项目设计上参考元服务轻量入口和快速触达思路，但当前主要按应用工程进行开发和调试。"],
            ["万能卡片", "可嵌入桌面或负一屏的轻量信息展示入口。", "项目配置了 WidgetCard，用于展示工地名称、温湿度、工人、监控和告警等摘要信息。"],
        ],
        [1.4, 2.4, 2.6],
    )

    doc.add_heading("3. 工程结构与 DevEco Studio 使用", level=1)
    doc.add_paragraph(
        "项目使用 DevEco Studio 进行开发、预览、构建和调试。工程目录中包含 AppScope、entry、hvigor、resources、ets 页面代码、桌面卡片代码以及构建配置文件。"
    )
    add_image(doc, PROJECT_STRUCTURE_IMAGE, "图 1  DevEco Studio 中的智慧工地工程结构", 6.2)
    add_table(
        doc,
        ["目录或文件", "作用"],
        [
            ["AppScope/app.json5", "配置应用级信息，例如 bundleName、版本号、应用图标和应用名称。"],
            ["entry/src/main/module.json5", "配置 entry 模块、UIAbility、权限、页面入口和 FormExtensionAbility。"],
            ["entry/src/main/ets/pages/Index.ets", "项目主页面，承担登录、看板、材料区、监控、打卡、工时、AI 助手等主要 UI 和交互逻辑。"],
            ["entry/src/main/ets/entryability/EntryAbility.ets", "应用主 Ability，负责加载 pages/Index 页面。"],
            ["entry/src/main/ets/entryformability/EntryFormAbility.ets", "桌面卡片扩展 Ability，负责卡片数据创建、刷新和事件处理。"],
            ["entry/src/main/ets/widget/pages/WidgetCard.ets", "万能卡片页面，使用 LocalStorageProp 接收卡片数据并展示。"],
            ["entry/src/main/resources", "应用图片、颜色、字符串、页面配置和卡片配置等资源。"],
            ["build-profile.json5 / hvigorfile.ts", "Hvigor 构建相关配置，用于打包 HAP。"],
        ],
        [2.1, 4.3],
    )

    doc.add_heading("4. ArkTS 与 ArkUI 声明式页面开发", level=1)
    doc.add_paragraph(
        "项目主页面采用 ArkTS 语言和 ArkUI 声明式范式开发。ArkUI 的特点是通过组件树描述界面，通过状态变量驱动界面刷新。项目中的登录页、管理员首页、材料区卡片、打卡按钮、底部导航栏、AI 悬浮按钮等都由 ArkUI 组件组合实现。"
    )
    add_table(
        doc,
        ["ArkUI 技术点", "项目中的实际使用"],
        [
            ["@Entry / @Component", "Index.ets 使用 @Entry 和 @Component 标识主页面组件。"],
            ["@State", "保存登录状态、当前角色、材料数据、打卡状态、当前选中 Tab、AI 助手开关等界面状态。"],
            ["Column / Row / Grid", "构建纵向布局、横向布局和今日看板的卡片网格。"],
            ["Button / Text / Image", "实现登录按钮、退出按钮、风扇 ON/OFF、AI 按钮、应用图标和数据展示。"],
            ["ForEach", "用于批量渲染装饰元素、数据列表、材料告警记录和数字孪生资产列表。"],
            ["animateTo", "用于登录、角色切换、页面切换、悬浮按钮打开等交互动画。"],
            ["Web 组件", "用于嵌入 THINGJS 数字孪生页面和云端视频播放页面。"],
        ],
        [1.6, 4.8],
    )
    doc.add_paragraph(
        "项目界面风格结合像素风与工地元素，使用黄色、黑色警示条、卡片边框、醒目的温湿度数字和底部导航来表达智慧工地主题。ArkUI 的链式样式写法使页面尺寸、颜色、圆角、阴影和点击事件可以集中写在组件后方，便于快速调整界面效果。"
    )
    add_image(doc, APP_HOME_IMAGE, "图 2  鸿蒙端智慧工地管理员首页运行效果", 3.0)

    doc.add_heading("5. UIAbility 生命周期与页面加载", level=1)
    doc.add_paragraph(
        "鸿蒙应用的页面入口由 UIAbility 管理。本项目的 EntryAbility 继承 UIAbility，在 onWindowStageCreate 生命周期中调用 windowStage.loadContent('pages/Index') 加载主页面。"
    )
    add_table(
        doc,
        ["生命周期方法", "项目中的作用"],
        [
            ["onCreate", "Ability 创建时输出日志，便于调试应用启动过程。"],
            ["onWindowStageCreate", "创建主窗口并加载 pages/Index，是应用进入主界面的关键步骤。"],
            ["onForeground", "应用进入前台时触发，可用于后续恢复刷新或重新拉取数据。"],
            ["onBackground", "应用进入后台时触发，可用于后续暂停计时或释放资源。"],
            ["onDestroy", "Ability 销毁时触发，用于日志记录和资源清理。"],
        ],
        [1.7, 4.7],
    )

    doc.add_heading("6. 权限配置与网络接口调用", level=1)
    doc.add_paragraph(
        "智慧工地项目需要访问华为云 IoTDA、Supabase、金仓数据库 API、Dify 接口、视频流页面和 THINGJS 页面，因此在 module.json5 中申请了 ohos.permission.INTERNET 权限。"
    )
    add_table(
        doc,
        ["接口方向", "鸿蒙端实现方式", "业务作用"],
        [
            ["华为云 IoTDA", "@ohos.net.http 发起请求", "读取设备影子，获取水泥区和木材区温湿度、风扇、灯等状态；下发风扇控制命令。"],
            ["Supabase 云数据库", "HTTP REST 请求", "读取打卡任务和签到记录，实现跨设备可见的打卡管理。"],
            ["金仓材料 API", "HTTP 请求本地或穿透接口", "读取材料历史数据，绘制材料区趋势图。"],
            ["Dify 大模型接口", "HTTP POST chat-messages", "将用户问题发送给智慧工地助手并接收回答。"],
            ["THINGJS / 视频页面", "Web 组件加载外部页面", "展示数字孪生和云端视频播放界面。"],
        ],
        [1.6, 2.1, 2.7],
    )
    doc.add_paragraph(
        "在代码结构上，项目将 Supabase、Dify、材料数据库和 IoTDA 的连接信息分别放入 common 目录中的配置文件，主页面再调用这些配置进行请求。这种方式便于后续修改接口地址、Token 或局域网 IP。"
    )

    doc.add_heading("7. 状态管理、定时刷新与实时数据显示", level=1)
    doc.add_paragraph(
        "项目使用 @State 保存界面状态，并结合定时器进行数据刷新。材料区数据每隔固定时间自动同步，打卡工时通过秒级定时器更新，用户不需要手动切换页面才能看到最新状态。"
    )
    for idx, item in enumerate(
        [
            "应用启动或进入页面后，初始化状态变量，例如登录状态、角色、当前 Tab 和材料区默认值。",
            "通过 setInterval 定时调用 refreshCloudData，从华为云 IoTDA 或后端接口刷新材料区数据。",
            "工人签到后启动工时计时逻辑，通过定时器按秒更新当前工时显示。",
            "用户点击风扇 ON/OFF 后，鸿蒙端向 IoTDA 发送命令，并根据请求结果更新提示状态。",
            "材料区点击后读取金仓数据库历史记录，在趋势图页面显示温湿度变化。",
        ],
        1,
    ):
        add_number(doc, idx, item)

    doc.add_heading("8. 桌面卡片开发", level=1)
    doc.add_paragraph(
        "项目不仅实现了主应用页面，还实现了鸿蒙桌面卡片。卡片由 FormExtensionAbility 和 WidgetCard 共同完成：EntryFormAbility 负责创建和更新卡片数据，WidgetCard 负责卡片 UI 展示。"
    )
    add_table(
        doc,
        ["卡片技术点", "项目中的实现"],
        [
            ["FormExtensionAbility", "EntryFormAbility 中实现 onAddForm、onUpdateForm、onFormEvent、onRemoveForm 等生命周期。"],
            ["formBindingData", "将工地名称、监控数量、温度、湿度、工人数、告警数和刷新时间绑定到卡片。"],
            ["LocalStorageProp", "WidgetCard 使用 LocalStorageProp 接收 FormExtensionAbility 传入的数据。"],
            ["postCardAction", "卡片点击可拉起 EntryAbility，也可以发送刷新事件。"],
            ["form_config.json", "配置卡片名称、尺寸、更新时间、默认规格和支持规格。"],
        ],
        [1.7, 4.7],
    )
    doc.add_paragraph(
        "桌面卡片的作用是让管理员不打开完整应用也能看到智慧工地摘要信息，例如温湿度、告警数量和刷新时间。这体现了 HarmonyOS 多入口、轻量触达的开发特点。"
    )

    doc.add_heading("9. Web 组件与跨技术融合", level=1)
    doc.add_paragraph(
        "HarmonyOS 应用并不只展示原生 ArkUI 页面，也可以通过 Web 组件接入外部可视化页面。本项目使用 Web 组件加载 THINGJS 数字孪生页面，并保留云端视频播放入口。"
    )
    add_table(
        doc,
        ["Web 场景", "项目用途"],
        [
            ["THINGJS 数字孪生", "展示工地与材料区 3D 场景，将材料区数据与空间展示结合。"],
            ["云端视频播放", "用于接入 RTSP 推流后的网页播放页面。"],
            ["外部接口页面", "便于在鸿蒙端嵌入已有 Web 可视化能力，减少重复开发成本。"],
        ],
        [1.8, 4.6],
    )

    doc.add_heading("10. 构建、调试与部署", level=1)
    doc.add_paragraph(
        "项目使用 Hvigor 进行构建。DevEco Studio 可以完成同步、编译、打包和真机运行。截图中的构建输出显示 entry:default@PackageHap、PackingCheck、SignHap、assembleHap 等任务已执行完成，并显示 BUILD SUCCESSFUL。"
    )
    add_image(doc, BUILD_SUCCESS_IMAGE, "图 3  DevEco Studio 中 HAP 构建成功结果", 6.2)
    add_table(
        doc,
        ["阶段", "说明"],
        [
            ["工程同步", "DevEco Studio 读取 oh-package、build-profile 和 hvigor 配置。"],
            ["ArkTS 编译", "将 Index.ets、Ability、WidgetCard 等 ArkTS 文件编译为可运行模块。"],
            ["资源处理", "打包图片、颜色、字符串、页面配置、卡片配置等 resources 内容。"],
            ["HAP 打包", "通过 assembleHap 生成可安装的 HarmonyOS 应用包。"],
            ["签名与安装", "调试阶段需要配置签名或使用调试签名，然后安装到模拟器或真机。"],
            ["真机调试", "通过 USB 调试连接 HarmonyOS 手机，运行应用并查看日志与界面。"],
        ],
        [1.5, 4.9],
    )

    doc.add_heading("11. 鸿蒙技术在项目中的综合价值", level=1)
    for item in [
        "移动端体验完整：通过 ArkUI 构建了管理员端、工人端、材料区、监控、打卡、工时和 AI 助手等页面。",
        "跨服务能力强：通过网络权限和 HTTP 能力接入华为云 IoTDA、Supabase、金仓数据库和 Dify。",
        "设备数据可视化：通过卡片、看板、趋势图和数字孪生，把工地数据以更直观的方式展示出来。",
        "多入口触达：主应用负责完整管理，桌面卡片负责快捷查看，AI 悬浮入口负责自然语言查询。",
        "适合实训展示：DevEco Studio 工程结构清晰，构建链路完整，能够体现鸿蒙应用从开发到运行的完整过程。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("12. 总结", level=1)
    doc.add_paragraph(
        "智慧工地系统在鸿蒙开发技术上的应用较为完整：使用 DevEco Studio 管理工程，使用 ArkTS 和 ArkUI 完成移动端界面，使用 UIAbility 加载主页面，使用 FormExtensionAbility 和 WidgetCard 实现桌面卡片，使用 INTERNET 权限和 HTTP 能力接入云端与本地服务，使用 Web 组件融合数字孪生与视频页面，并通过 Hvigor 完成 HAP 构建。"
    )
    doc.add_paragraph(
        "从项目实现角度看，鸿蒙技术不仅承担了页面显示，还承担了移动端业务编排、设备状态展示、远程接口调用、桌面卡片触达和跨平台页面融合等关键任务。该项目能够较好体现 HarmonyOS 应用开发在智慧工地场景中的实际应用价值。"
    )

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_font(run)

    doc.core_properties.title = "智慧工地系统鸿蒙开发技术应用说明"
    doc.core_properties.subject = "HarmonyOS 智慧工地项目技术应用总结"
    doc.core_properties.author = "Codex"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
