# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "personal_report_huawei_work.docx"
OUTPUT_DOCX = Path(r"C:\Users\lenovo\Desktop\软件系统实训报告_个人完成版_华为云补充.docx")
ASSET_DIR = ROOT / "docs" / "personal_report_huawei_assets"
ARCH_IMAGE = ASSET_DIR / "system_architecture_huawei.png"
CONSOLE_IMAGE = ASSET_DIR / "huawei_iotda_console.png"


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def insert_paragraph_before(paragraph, text, style_name, size=10.5, bold=False):
    new_para = paragraph.insert_paragraph_before()
    try:
        new_para.style = style_name
    except Exception:
        pass
    run = new_para.add_run(text)
    set_font(run, size=size, bold=bold)
    if style_name == "Body Text First Indent 2":
        new_para.paragraph_format.line_spacing = 1.25
    return new_para


def insert_picture_before(paragraph, image_path, caption, width):
    picture_para = paragraph.insert_paragraph_before()
    picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_para.add_run()
    run.add_picture(str(image_path), width=Inches(width))

    caption_para = paragraph.insert_paragraph_before()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_para.add_run(caption)
    set_font(run, size=9, color="666666")


def main():
    doc = Document(DOCX_PATH)
    target = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("3.5") and "Testing" in paragraph.text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError("未找到 3.5 测试Testing 插入位置")

    insert_paragraph_before(target, "华为云 IoTDA 联动模块补充", "Heading 2", size=11, bold=True)
    insert_paragraph_before(
        target,
        "在项目整体架构中，华为云 IoTDA 是连接硬件设备与鸿蒙应用的关键云平台。板端传感器负责采集温湿度、烟雾、CO2、风扇等现场数据，数据上报到华为云 IoTDA 后，鸿蒙端通过设备影子接口读取最新状态，并在管理员端材料区看板中实时展示。该联动使系统不只是静态页面，而是能够接入真实设备数据的智慧工地应用。",
        "Body Text First Indent 2",
    )
    insert_picture_before(target, ARCH_IMAGE, "图 9  智慧工地系统中华为云 IoTDA 联动架构", 6.4)

    insert_paragraph_before(
        target,
        "本人在华为云控制台中配置 IoTDA 实例、产品、设备和设备影子访问参数，并在 DevEco Studio 项目中维护 projectId、deviceId、设备影子 URL、Token 与命令下发参数。鸿蒙端通过定时刷新机制读取云端影子数据，使水泥区和木材区温湿度能够同步到页面；管理员点击风扇 ON/OFF 后，系统通过 IoTDA 命令下发接口将控制指令投递到设备侧。",
        "Body Text First Indent 2",
    )
    insert_picture_before(target, CONSOLE_IMAGE, "图 10  华为云 IoTDA 实例与设备接入控制台", 6.4)

    insert_paragraph_before(
        target,
        "在调试过程中，华为云联动主要涉及三类问题：第一是 Token、projectId、deviceId 等鉴权和资源标识必须匹配；第二是设备影子读取成功后，ArkUI 状态变量需要及时更新，否则会出现“日志有数据但卡片不显示”的问题；第三是命令下发成功并不等于设备状态立即变化，板端必须正确订阅命令 topic、执行控制逻辑并重新上报属性，应用端才能在下一轮影子同步中看到状态变化。",
        "Body Text First Indent 2",
    )
    insert_paragraph_before(
        target,
        "通过该模块，系统实现了“传感器/执行器—华为云 IoTDA—鸿蒙应用—数据库/数字孪生”的数据闭环：实时数据用于看板展示，历史数据写入金仓数据库用于趋势分析，设备状态和告警信息又可以进一步服务数字孪生场景和 AI 助手问答。",
        "Body Text First Indent 2",
    )

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
