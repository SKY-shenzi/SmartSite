# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "docs" / "training_report_template.docx"
OUTPUT = Path(r"C:\Users\lenovo\Desktop\软件系统实训报告_个人完成版.docx")
ASSET_DIR = ROOT / "docs" / "training_report_assets"

IMAGE_MAP = {
    "home": ASSET_DIR / "HarmonyOS主页面.png",
    "ai_page": ASSET_DIR / "HarmonyAI助手.png",
    "ai_flow": ASSET_DIR / "AI工作流编排.png",
    "twin_page": ASSET_DIR / "HarmonyOS数字孪生环节.png",
    "thingjs": ASSET_DIR / "数字孪生具体场景搭建.png",
    "supabase": ASSET_DIR / "superbase数据库设计图.png",
    "kingbase": ASSET_DIR / "金仓数据库启动图片.png",
    "video": ASSET_DIR / "视频流传输图片.png",
}


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_paragraph(paragraph):
    paragraph.text = ""


def set_para_text(paragraph, text, size=10.5, bold=False):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_para(doc, text, style="Body Text First Indent 2"):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    paragraph.paragraph_format.line_spacing = 1.25
    return paragraph


def add_heading(doc, text, level):
    style = "Heading 3" if level == 1 else "Heading 2"
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=12 if level == 1 else 11, bold=True)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="Body Text First Indent 2")
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.left_indent = Inches(0.22)
    run = paragraph.add_run("（1）" + text if False else "• " + text)
    set_run_font(run, size=10.5)
    return paragraph


def add_numbered(doc, number, text):
    paragraph = doc.add_paragraph(style="Body Text First Indent 2")
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.left_indent = Inches(0.20)
    run = paragraph.add_run(f"{number}. {text}")
    set_run_font(run, size=10.5)
    return paragraph


def add_image(doc, key, caption, width):
    image_path = IMAGE_MAP[key]
    if not image_path.exists():
        return
    doc.add_picture(str(image_path), width=Inches(width))
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_para.add_run(caption)
    set_run_font(run, size=9, color="666666")


def remove_body_after_cover(doc):
    # Keep the original cover/title/table and rebuild body content using template styles.
    # Paragraph 13 is the first body heading in the converted template.
    body = doc._body._element
    paragraphs = list(doc.paragraphs)
    for paragraph in paragraphs[13:]:
        body.remove(paragraph._element)


def fill_cover(doc):
    table = doc.tables[0]
    # Leave personal fields empty for the user to fill; fill known report metadata only.
    table.cell(4, 1).text = "2026年7月"
    table.cell(5, 1).text = "2026年7月11日"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10.5)


def add_code_block(doc, lines):
    for line in lines:
        paragraph = doc.add_paragraph(style="Body Text First Indent 2")
        paragraph.paragraph_format.first_line_indent = Inches(0)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)


def build_report():
    doc = Document(TEMPLATE)
    fill_cover(doc)
    remove_body_after_cover(doc)

    add_heading(doc, "实践目的 Aim of Project", 1)
    add_para(
        doc,
        "本次软件系统开发实践以“智慧工地项目”为主题，目标是在真实项目场景下完成一个具备移动端交互、云端设备接入、数据库存储、视频监控、数字孪生和大模型助手能力的软件系统。通过该项目，本人希望系统掌握 HarmonyOS 应用开发、云平台接口调用、国产数据库应用、前后端数据交互和项目综合集成的方法。"
    )
    add_para(
        doc,
        "实践过程中，本人作为项目主要推进者，围绕智慧工地的管理员端、工人端、材料区监测、打卡签到、华为云 IoTDA 数据同步、风扇命令下发、Supabase 云数据库、金仓数据库、Dify 大模型助手、THINGJS 数字孪生和 RTSP 视频流接入等内容开展开发与调试。项目最终形成了一个可运行、可展示、可说明数据来源的综合型实训系统。"
    )

    add_heading(doc, "预习内容 Preview Content", 1)
    add_para(
        doc,
        "在正式开发前，本人围绕项目所需技术进行了预习和资料准备。首先学习 HarmonyOS 应用、元服务和万能卡片的基本概念，理解 DevEco Studio 工程结构、ArkTS 语法、ArkUI 声明式组件、UIAbility 生命周期、FormExtensionAbility 卡片能力以及网络权限配置。其次学习云端与数据库相关知识，包括华为云 IoTDA 设备影子、命令下发、IAM Token、Supabase REST API、金仓数据库 KingbaseES 建表和 Node.js 后端接口。"
    )
    add_para(
        doc,
        "此外，本人还预习了视频流与数字孪生相关技术，包括 RTSP/RTMP/HLS 视频流传输、SRS 服务、摄像头推流、THINGJS 场景搭建、3D 场景分享链接嵌入，以及 Dify + Ollama + DeepSeek 本地大模型工作流编排。通过这些预习内容，项目开发从单一移动端页面扩展为一个包含云、端、数据库和 AI 的综合系统。"
    )

    add_heading(doc, "实践内容和实践过程 Project Content and Process", 1)

    add_heading(doc, "3.1\t概述Overview", 2)
    add_para(
        doc,
        "智慧工地项目面向施工现场管理场景，系统分为工人端和管理员端。工人端主要负责打卡签到、签退和秒级工时显示；管理员端主要负责查看今日看板、材料区监测、视频监控、打卡任务管理、工时查看、数字孪生展示和 AI 助手问答。"
    )
    add_para(
        doc,
        "本人承担的主要工作包括：项目前期需求规划、页面结构设计、HarmonyOS 主页面开发、工人端与管理员端逻辑区分、登录与退出逻辑修正、材料区华为云数据接入、风扇命令下发、打卡系统云数据库接入、金仓数据库材料历史数据功能、大模型助手接入、数字孪生场景嵌入、视频流接入规划与展示，以及多轮界面优化和问题排查。"
    )
    add_image(doc, "home", "图 1  HarmonyOS 智慧工地管理员端主页面", 3.1)

    add_heading(doc, "3.2\t相关技术Relevant Technologies", 2)
    add_para(
        doc,
        "本项目使用的核心技术覆盖移动端、云平台、数据库、视频流、数字孪生和人工智能多个方向。移动端部分采用 HarmonyOS SDK、ArkTS 和 ArkUI 进行开发，在 DevEco Studio 中完成工程管理、编译、打包和调试。主页面使用 @Entry、@Component、@State、Column、Row、Grid、Button、Text、Image、Web 等 ArkUI 能力实现。"
    )
    add_para(
        doc,
        "云平台方面，项目接入华为云 IoTDA。材料区设备通过设备影子上报温度、湿度、灯光、风扇等状态，鸿蒙端定时读取影子数据并展示；管理员点击风扇 ON/OFF 后，系统向华为云下发命令，再由云端投递到板端设备。"
    )
    add_para(
        doc,
        "数据库方面，打卡签到模块使用 Supabase 云数据库，便于多台设备共享打卡任务和记录；材料监测历史模块使用国产金仓数据库 KingbaseES，通过 Node.js 后端 API 保存华为云同步来的材料区历史数据，并在鸿蒙端绘制折线趋势图。"
    )
    add_para(
        doc,
        "人工智能方面，系统接入 Dify 工作流和 Ollama 本地模型，构建“智慧工地助手”。助手能够根据用户问题区分材料监测、打卡签到、视频监控、数字孪生和项目说明等类型，通过 HTTP 节点读取项目数据，再由 DeepSeek 模型生成自然语言回答。"
    )
    add_image(doc, "ai_page", "图 2  HarmonyOS 端智慧工地助手页面", 3.0)
    add_image(doc, "ai_flow", "图 3  Dify 大模型工作流编排", 6.2)

    add_heading(doc, "3.3\t系统设计System Design", 2)
    add_para(
        doc,
        "系统总体采用“鸿蒙移动端 + 云平台 + 数据库 + 后端接口 + AI 工作流 + 数字孪生页面”的结构。鸿蒙端负责用户交互和业务展示；华为云 IoTDA 负责设备数据与命令下发；Supabase 负责打卡相关云数据；金仓数据库负责材料监测历史数据；Node.js API 作为鸿蒙端与金仓数据库之间的中间层；Dify 负责 AI 助手工作流编排；THINGJS 负责数字孪生场景展示。"
    )
    add_para(
        doc,
        "在角色设计上，系统区分工人端与管理员端。工人端只保留打卡签到核心页面，若没有打卡任务则显示暂无任务；管理员端提供今日看板、材料、监控、打卡和工时等导航。该设计避免了工人端误切换到管理员端的逻辑问题，也让两个端的职责更加清晰。"
    )
    add_para(
        doc,
        "材料区设计上，项目设置水泥区和木材区两个差异较大的存储场景。水泥区和木材区拥有不同的存储阈值、负责人、设备 ID 和控制命令格式。材料区卡片展示温度、湿度、风扇、告警和灯光状态，点击材料区可进入历史趋势页面，从金仓数据库读取并显示折线统计图。"
    )
    add_para(
        doc,
        "数据库设计上，Supabase 用于打卡系统，包含用户、打卡任务和打卡记录等表；金仓数据库用于材料监测，设计 users、material_zones、material_sensor_records、material_alert_rules、material_alert_events、digital_twin_assets 等 6 张表，满足国产数据库建立 3 张以上数据表并完成业务功能的要求。"
    )
    add_image(doc, "supabase", "图 4  Supabase 云数据库表设计", 6.2)
    add_image(doc, "kingbase", "图 5  金仓材料 API 启动界面", 6.0)
    add_para(
        doc,
        "数字孪生设计上，本人使用 THINGJS 搭建材料区和工地空间场景，并将其通过 WebView 嵌入鸿蒙应用。数字孪生场景用于展示材料仓储区、传感器、风扇、灯光等对象，使用户不仅能看到数据，也能在 3D 场景中理解工地空间关系。"
    )
    add_image(doc, "thingjs", "图 6  THINGJS 数字孪生场景搭建", 6.2)
    add_image(doc, "twin_page", "图 7  HarmonyOS 端数字孪生展示页面", 3.0)

    add_heading(doc, "3.4\t编码Coding", 2)
    add_para(
        doc,
        "编码阶段，本人主要围绕 HarmonyOS 主应用、接口配置、云端数据同步、数据库交互和 AI 助手调用进行实现。项目主页面集中在 entry/src/main/ets/pages/Index.ets 中，配置文件包括 IotdaConfig.ets、SupabaseConfig.ets、MaterialDatabaseConfig.ets 和 DifyConfig.ets。"
    )
    add_para(
        doc,
        "在 HarmonyOS 页面编码方面，本人使用 ArkUI 构建像素风与工地元素结合的界面，包括顶部标题栏、今日看板、材料区卡片、底部导航栏、AI 悬浮按钮、打卡圆形按钮、历史趋势页面和 WebView 页面。为了提升体验，在角色切换、登录、页面切换和 AI 按钮打开时加入 animateTo 动画。"
    )
    add_para(
        doc,
        "在打卡系统编码方面，本人实现了管理员发布打卡任务、关闭任务、重置记录、查看打卡情况，工人端实现单按钮签到/签退和秒级工时显示。打卡记录写入 Supabase 云数据库，使不同手机或不同端可以访问同一份打卡任务和记录。"
    )
    add_para(
        doc,
        "在华为云 IoTDA 编码方面，本人实现了设备影子查询和风扇命令下发。系统每隔固定时间读取水泥区和木材区设备影子，解析 Temp、Humi、LampSt、FanSt 等属性并显示在材料区。风扇控制根据不同产品模型使用不同命令：水泥区使用 SetFan/FanSt，木材区使用 Fan/fan，避免命令参数不匹配导致下发失败。"
    )
    add_para(
        doc,
        "在金仓数据库编码方面，本人将移动端与数据库解耦，使用 Node.js 实现 Kingbase Material API，端口为 8088。鸿蒙端通过 HTTP 请求向 /api/material-records 写入材料数据，通过 /api/material-records?zoneCode=...&limit=... 查询历史记录，并绘制温湿度折线趋势。"
    )
    add_code_block(
        doc,
        [
            "金仓材料 API 启动命令：",
            "cd D:\\SmartSiteServices\\kingbase-material-api",
            "cmd /c npm start",
            "材料历史查询示例：",
            "GET http://电脑IP:8088/api/material-records?zoneCode=cement_zone&limit=24",
        ],
    )
    add_para(
        doc,
        "在 AI 助手编码方面，本人通过 DifyConfig.ets 配置 Dify chat-messages 接口，鸿蒙端把用户输入作为 query 发送给 Dify。Dify 工作流内部使用问题分类器、HTTP 请求、知识检索、代码执行和 LLM 节点完成问答。该功能让用户可以询问打卡、材料监测、视频监控、数字孪生和项目运维相关问题。"
    )
    add_para(
        doc,
        "在视频流编码与规划方面，本人围绕 RTSP 推流和云服务器播放进行设计。项目中保留视频监控入口，支持通过 Web 页面加载视频流播放地址，同时结合 SRS、RTSP/RTMP/HLS 等方案进行调试说明。"
    )
    add_image(doc, "video", "图 8  视频流传输与播放链路", 6.2)

    add_heading(doc, "3.5\t测试Testing", 2)
    add_para(
        doc,
        "测试阶段，本人对项目进行了多轮功能测试、接口测试、真机/预览测试和异常排查。首先在 DevEco Studio 中进行构建测试，确认 HAP 能够成功生成，解决 compatibleSdkVersion、releaseType、设备检测和签名相关问题。"
    )
    add_para(
        doc,
        "其次对华为云 IoTDA 接入进行测试，先通过 API Explorer 和 Token 调试确认设备影子 URL、projectId、deviceId、实例 ID 等参数，再在鸿蒙端测试材料区数据读取。实际调试中曾出现“能读到数据但卡片不显示”的问题，后续通过状态更新和页面刷新逻辑修复，使温湿度值能够正常进入 UI。"
    )
    add_para(
        doc,
        "对命令下发功能，本人通过华为云消息跟踪确认命令已经投递到设备 topic，并根据云端错误提示修正木材区命令名称和参数。水泥区和木材区命令模型不同，因此分别适配不同 command_name 和 paras，最终实现风扇 ON/OFF 控制请求能够成功下发。"
    )
    add_para(
        doc,
        "对打卡系统，本人测试了管理员发布任务、关闭任务、重置记录、工人签到、签退和秒级工时计算。针对工人端曾经能切换管理员端、工时不计时、签到签退按钮割裂等问题，后续将其调整为单圆形按钮并按秒显示工时。"
    )
    add_para(
        doc,
        "对数据库部分，本人测试了 Supabase 用户和打卡表访问，也测试了金仓材料 API 的启动、材料数据写入和历史查询。金仓部分使用接口层访问数据库，避免鸿蒙端直接保存数据库密码。"
    )
    add_para(
        doc,
        "对数字孪生和视频流部分，本人分别测试了 THINGJS 场景分享链接在鸿蒙 WebView 中的加载效果，以及 RTSP 视频流地址和网页端播放方式。针对数字孪生页面黑屏、闪烁、按钮突兀等问题，进行了多轮链接替换和界面调整。"
    )

    add_heading(doc, "实践总结 Conclusion", 1)
    add_para(
        doc,
        "通过本次实训，本人从一个单纯的鸿蒙案例开发逐步推进到智慧工地综合系统开发，实践内容覆盖移动端、云平台、数据库、AI、视频流和数字孪生多个方向。项目过程中最大的收获是理解了一个真实软件系统不是单个页面或单个接口，而是由角色权限、数据来源、业务流程、界面体验、后端服务和调试部署共同组成。"
    )
    add_para(
        doc,
        "本次项目也暴露出许多实际开发问题，例如 Token 失效、设备影子和命令下发状态不一致、接口地址在手机和电脑之间不同、数据库服务需要先启动、视频流本机地址不能被其他设备访问、Dify 节点能拿到数据但不会自动返回给 LLM 等。这些问题都不是单看代码就能解决的，需要结合日志、接口返回、云平台消息跟踪和实际运行环境综合排查。"
    )
    add_para(
        doc,
        "总体来看，本人完成了智慧工地项目的主要功能搭建和多项关键集成工作，包括 HarmonyOS 应用界面、双端角色逻辑、打卡系统、材料监测、华为云 IoTDA、风扇命令下发、Supabase、金仓数据库、Dify 大模型助手、THINGJS 数字孪生和视频流接入。后续如果继续完善，可以进一步加强权限安全、视频智能识别、数字孪生实时动画绑定和数据库完整 CRUD 管理。"
    )

    add_heading(doc, "参考资料 References", 1)
    refs = [
        "[1] 华为开发者联盟. HarmonyOS 应用开发文档、ArkTS 与 ArkUI 开发指南.",
        "[2] 华为云. IoTDA 设备接入服务 API 文档、设备影子与命令下发说明.",
        "[3] Supabase. REST API、Authentication 与数据库表设计相关文档.",
        "[4] 金仓数据库 KingbaseES 官方文档与数据库安装、建表、连接说明.",
        "[5] Dify 官方文档. Workflow、HTTP 请求节点、知识库与模型供应商配置说明.",
        "[6] Ollama 官方文档. 本地模型部署、模型拉取与服务启动说明.",
        "[7] THINGJS 官方文档. 场景搭建、分享链接与数字孪生开发说明.",
        "[8] SRS / FFmpeg / RTSP 相关资料. 视频流推送、转码与播放说明.",
    ]
    for ref in refs:
        add_para(doc, ref)

    # Normalize fonts in all generated runs.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name != "Consolas":
                set_run_font(run)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
